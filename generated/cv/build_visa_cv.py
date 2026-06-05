import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "generated" / "cv"
NODE = "/Users/enya/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"


def load_site_data():
    js = """
const fs = require('fs');
const vm = require('vm');
const code = fs.readFileSync('assets/js/data.js', 'utf8');
const sandbox = { window: {} };
vm.createContext(sandbox);
vm.runInContext(code, sandbox);
process.stdout.write(JSON.stringify(sandbox.window.SITE_DATA));
"""
    raw = subprocess.check_output([NODE, "-e", js], cwd=ROOT)
    return json.loads(raw)


def strip_html(text):
    text = re.sub(r"</?(?:b|strong)>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.replace("&amp;", "&").strip()


def pick(value, lang):
    if isinstance(value, dict) and "zh" in value:
        return value.get(lang) or value.get("zh") or ""
    return value or ""


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_style_font(style, name, size, color=None, bold=False):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def setup_doc(lang):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    base_font = "Microsoft YaHei" if lang == "zh" else "Calibri"
    set_style_font(doc.styles["Normal"], base_font, 10.2, "222222")
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.08
    for style_name, size, color in [
        ("Heading 1", 13.5, "2E74B5"),
        ("Heading 2", 11.5, "2E74B5"),
        ("Heading 3", 10.8, "1F4D78"),
    ]:
        set_style_font(doc.styles[style_name], base_font, size, color, True)
        doc.styles[style_name].paragraph_format.space_before = Pt(8)
        doc.styles[style_name].paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    return doc


def add_title(doc, title, subtitle, email, lang):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = "Microsoft YaHei" if lang == "zh" else "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(subtitle)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("555555")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(email)
    run.font.size = Pt(10)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, [1900, 7400])
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.6)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


def add_section(doc, title):
    doc.add_heading(title, level=1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_timeline(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, [1800, 7500])
    set_table_borders(table, "E5E7EB")
    for period, text in rows:
        row = table.add_row()
        row.cells[0].text = period
        row.cells[1].text = text
        set_cell_shading(row.cells[0], "F7F9FB")
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9.4)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


PUB_EN_OVERRIDES = {
    "典型三维模型表示方法转换技术综述": {
        "authors": "Yin Xiaolong, Enya Shen*, Wang Jianmin",
        "title": "A Review of Conversion Techniques for Typical 3D Model Representation Methods",
        "venue": "Journal of Computer-Aided Design & Computer Graphics 37(10), 2025: 1666-1680",
    },
    "B-Rep 模型高效布尔运算算法": {
        "authors": "Liang Sili, Huang Haomian, Liang Hao, Tang Weiran, Bao Anchang, Enya Shen*, Wang Jianmin, Yang Yijun",
        "title": "Efficient Boolean Operations for B-Rep Models",
        "venue": "Journal of Computer-Aided Design & Computer Graphics 37(11), 2025: 1993-2005",
    },
    "三维船舶模型边界表示转构造表示方法": {
        "authors": "Tai Zhiwei, Zhang Xinglong, Yin Xiaolong, Enya Shen*, Wang Jianmin",
        "title": "A Method for Transforming B-Rep of 3D Ship Models into CSG",
        "venue": "Journal of Computer-Aided Design & Computer Graphics 35(12), 2023: 1851-1862",
    },
    "大数据可视化技术及应用": {
        "authors": "Enya Shen",
        "title": "Big Data Visualization Techniques and Applications",
        "venue": "Science & Technology Review 38(3), 2020: 68-83",
    },
    "大规模数据并行可视化与交互环境": {
        "authors": "Enya Shen, Wang Pan, Li Sikun, Cai Xun, Zeng Liang, Wang Wenke",
        "title": "A Parallel Visualization and Interaction Environment for Large-Scale Data",
        "venue": "HPC China 2012 (Best Paper Nomination)",
    },
}


PATENT_EN_OVERRIDES = {
    "一种基于样例的可视化生成方法及系统": {
        "authors": "Wang Jianmin, Shen Enya, Shen Leixian, Tai Zhiwei",
        "title": "Example-Based Visualization Generation Method and System",
    },
    "债券产品的潜在客户推荐方法和装置": {
        "authors": "Mao Guanzhong, Zuo Guangyuan, Wang Jianmin, Zhao Xibin, Shen Enya, Song Yiran",
        "title": "Method and Apparatus for Recommending Potential Customers for Bond Products",
    },
    "图布局优化方法和系统、电子设备及可读存储介质": {
        "authors": "Wang Jianmin, Shen Enya, Tai Zhiwei, Song Yiran, Shen Leixian",
        "title": "Graph Layout Optimization Method and System, Electronic Device, and Readable Storage Medium",
    },
    "一种自适应采样与查询方法及系统": {
        "authors": "Wang Jianmin, Shen Enya, Song Yiran, Shen Leixian",
        "title": "Adaptive Sampling and Query Method and System",
    },
    "一种社交网络关键节点发现方法及系统": {
        "authors": "Wang Jianmin, Shen Enya, Tai Zhiwei, Song Yiran",
        "title": "Social Network Key Node Discovery Method and System",
    },
    "一种面向任务的可视化推荐方法及装置": {
        "authors": "Wang Jianmin, Shen Leixian, Shen Enya, Tai Zhiwei, Song Yiran",
        "title": "Task-Oriented Visualization Recommendation Method and Device",
    },
    "一种多核心元素的知识图谱边图布局方法及系统": {
        "authors": "Wang Jianmin, Shen Enya, Zhang Li, Qian Zhou",
        "title": "Knowledge Graph Edge Layout Method and System Based on Multiple Core Elements",
    },
    "一种面向知识图谱的交互可视化方法、平台和系统": {
        "authors": "Wang Jianmin, Shen Enya, Zhang Li, Wei Yuhua",
        "title": "Interactive Visualization Method, Platform, and System for Knowledge Graphs",
    },
    "一种基于球面网格遥感数据文件集的键值存储方法及装置": {
        "authors": "Wang Jianmin, Shen Enya, Liu Yingbo, Zhao Xin",
        "title": "Key-Value Storage Method and Device for Spherical-Grid Remote-Sensing Data File Sets",
    },
    "一种基于模糊测度的流场涡特征检测方法": {
        "authors": "Zeng Liang, Xu Huaxun, Wang Huaihui, Cai Xun, Li Sikun, Wang Wenke, Shen Enya",
        "title": "Flow-Field Vortex Feature Detection Method Based on Fuzzy Measures",
    },
    "一种渐进最优的三角网格重复顶点快速去除方法": {
        "authors": "Wang Wenke, Wang Pan, Cai Xun, Li Sikun, Zeng Liang, Shen Enya, Wang Huaihui, Liu Huahai",
        "title": "Progressively Optimal Fast Duplicate-Vertex Removal Method for Triangular Meshes",
    },
    "基于双控制体的格心网格数据三维激波特征定位方法": {
        "authors": "Cai Xun, Ma Qianli, Wang Wenke, Zeng Liang, Li Sikun, Shen Enya, Wang Pan, Wang Wentao",
        "title": "Three-Dimensional Shock Feature Localization Method for Cell-Centered Grid Data Based on Dual Control Volumes",
    },
    "基于旋转因子的曲面流场涡特征提取方法": {
        "authors": "Li Sikun, Wang Huaihui, Zeng Liang, Cai Xun, Wang Wenke, Xu Huaxun, Wang Pan, Shen Enya",
        "title": "Curved-Surface Flow-Field Vortex Feature Extraction Method Based on Rotation Factors",
    },
    "基于 GPU 的三维空间交互点判定加速方法": {
        "authors": "Li Sikun, Shen Enya, Cai Xun, Wang Wenke, Zeng Liang, Wang Pan, Wang Huaihui, Wang Wentao",
        "title": "GPU-Based Acceleration Method for Determining Interactive Points in Three-Dimensional Space",
    },
}


def publication_text(pub, lang):
    if lang == "en" and pub["title"] in PUB_EN_OVERRIDES:
        item = PUB_EN_OVERRIDES[pub["title"]]
        return f"{item['authors']}. {item['title']}. {item['venue']}."
    return f"{strip_html(pub['authors'])}. {pub['title']}. {pub['venue']}."


def patent_text(pat, lang):
    if lang == "en" and pat["title"] in PATENT_EN_OVERRIDES:
        item = PATENT_EN_OVERRIDES[pat["title"]]
        return f"{item['authors']}. {item['title']}. {pat['id']}, {pat['date']}."
    return f"{pat['authors']}. {pat['title']}. {pat['id']}, {pat['date']}."


def build_cv(lang, data):
    zh = lang == "zh"
    doc = setup_doc(lang)
    name = pick(data["profile"]["name"], lang)
    title = "个人简历" if zh else "Curriculum Vitae"
    subtitle = (
        f"{pick(data['profile']['title'], lang)} | {pick(data['profile']['affiliation'], lang)}"
    )
    add_title(doc, f"{name} - {title}" if zh else f"{name} - {title}", subtitle, data["profile"]["email"], lang)

    add_section(doc, "个人信息" if zh else "Personal Information")
    add_kv_table(
        doc,
        [
            ("姓名" if zh else "Name", name),
            ("现职" if zh else "Current Position", pick(data["profile"]["title"], lang)),
            ("单位" if zh else "Affiliation", pick(data["profile"]["affiliation"], lang)),
            ("电子邮箱" if zh else "Email", data["profile"]["email"]),
            ("研究方向" if zh else "Research Areas", "、".join(pick(data["profile"]["areas"], lang)) if zh else ", ".join(pick(data["profile"]["areas"], lang))),
            ("更新日期" if zh else "Updated", "2026 年 6 月" if zh else "June 2026"),
        ],
    )

    add_section(doc, "个人简介" if zh else "Profile")
    p = doc.add_paragraph()
    p.add_run(strip_html(pick(data["profile"]["bio"], lang)))

    add_section(doc, "工作经历" if zh else "Employment")
    work_rows = [(item["period"], item["text"]) for item in pick(data["experience"]["work"], lang)]
    add_timeline(doc, work_rows)

    add_section(doc, "教育背景" if zh else "Education")
    edu_rows = [(item["period"], item["text"]) for item in pick(data["experience"]["education"], lang)]
    add_timeline(doc, edu_rows)

    add_section(doc, "研究方向" if zh else "Research Interests")
    add_bullets(
        doc,
        [
            f"{pick(area['title'], lang)}: {pick(area['desc'], lang)}"
            for area in data["research"]["areas"]
        ],
    )

    add_section(doc, "主要科研项目" if zh else "Selected Research Projects")
    add_bullets(
        doc,
        [
            f"{proj['period']}: {pick({'zh': proj['zh'], 'en': proj['en']}, lang)}"
            for proj in data["projects"]
        ],
    )

    add_section(doc, "论文发表" if zh else "Publications")
    add_numbered(doc, [publication_text(pub, lang) for pub in data["publications"]])

    add_section(doc, "专著" if zh else "Books")
    add_bullets(doc, [pick(book, lang) for book in data["books"]])

    add_section(doc, "授权/公开专利" if zh else "Patents")
    add_numbered(doc, [patent_text(pat, lang) for pat in data["patents"]])

    add_section(doc, "学生指导" if zh else "Student Mentoring")
    mentoring = []
    for group in ("current", "alumni"):
        for key, value in data["students"][group].items():
            label = pick(value["label"], lang)
            for item in value["items"]:
                mentoring.append(f"{label}, {item['year']}: {pick(item['names'], lang)}")
    add_bullets(doc, mentoring)

    add_section(doc, "服务与荣誉" if zh else "Service and Honors")
    service = pick(data["service"], lang)
    honors = [pick(h, lang) for h in data["honors"]]
    add_bullets(doc, service + honors)

    filename = "沈恩亚_美签个人简历_中文.docx" if zh else "Enya_Shen_US_Visa_CV_English.docx"
    path = OUT / filename
    doc.save(path)
    return path


def main():
    data = load_site_data()
    zh = build_cv("zh", data)
    en = build_cv("en", data)
    print(zh)
    print(en)


if __name__ == "__main__":
    main()
